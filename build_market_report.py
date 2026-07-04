from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ──────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x21, 0x37)
GREEN  = RGBColor(0x2A, 0x60, 0x41)
GOLD   = RGBColor(0xC9, 0x96, 0x3B)
DARK   = RGBColor(0x33, 0x33, 0x33)
GRAY   = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

BG_GREEN = "EBF4EE"
BG_GOLD  = "FDF6E8"
BG_NAV   = "0D2137"
BG_GRN   = "2A6041"
BORDER_C = "CCCCCC"

# ── Helpers ────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color=BORDER_C):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        margins.append(el)
    tcPr.append(margins)

def cell_paragraph(cell, text, bold=False, size=10, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Arial'
    return para

def add_cell_para(cell, text, bold=False, size=10, color=DARK, space_before=2):
    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Arial'
    return para

def style_header_row(row, bg=BG_NAV, text_color=WHITE):
    for cell in row.cells:
        set_cell_bg(cell, bg)
        set_cell_borders(cell)
        set_cell_margins(cell)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = text_color
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)

def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = 'Arial'
    if level == 1:
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
        # add bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '2A6041')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = GREEN
    elif level == 3:
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
    return para

def add_body(doc, text, italic=False, size=10.5, color=DARK):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return para

def add_bullet(doc, text, bold=False):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.bold = bold
    return para

def add_caveat(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(f"⚠  {text}")
    run.font.name = 'Arial'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)
    run.italic = True
    return para

def add_spacer(doc, size_pt=8):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(size_pt)
    return para

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_breaks.WD_BREAK.PAGE)

# Use run.add_break properly
from docx.enum.text import WD_BREAK
import docx.enum.text as docx_breaks

def set_col_widths(table, widths_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[i])

def make_highlight_box(doc, title, body_text, bg=BG_GREEN, border_color="2A6041"):
    """Single-cell table as a styled callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, border_color)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY; r.font.name = 'Arial'
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body_text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK; r2.font.name = 'Arial'
    return table

# ──────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ──────────────────────────────────────────────────────────────────

doc = Document()

# Page setup: US Letter, 1" margins
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)

# Default paragraph style
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)

# ── COVER PAGE ─────────────────────────────────────────────────────

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.space_after = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FAIRWAY GOLF CLUB")
r.font.name = 'Arial'; r.bold = True; r.font.size = Pt(32); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(8)
r2 = p2.add_run("LOCAL MARKET ANALYSIS")
r2.font.name = 'Arial'; r2.bold = True; r2.font.size = Pt(22); r2.font.color.rgb = GREEN

# Divider line
p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_div.paragraph_format.space_before = Pt(4); p_div.paragraph_format.space_after = Pt(20)
pPr = p_div._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12'); bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), 'C9963B')
pBdr.append(bot); pPr.append(pBdr)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(50)
r3 = p3.add_run("Investment Opportunity Assessment")
r3.font.name = 'Arial'; r3.italic = True; r3.font.size = Pt(14); r3.font.color.rgb = GRAY

for label, text in [
    ("Proposed Location:", "Workhouse Arts Center"),
    ("Address:", "9518 Workhouse Way, Lorton, Virginia 22079"),
    ("Date:", "July 2026"),
    ("Classification:", "CONFIDENTIAL — Prepared for investor use only"),
]:
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(2); pp.paragraph_format.space_after = Pt(2)
    rl = pp.add_run(label + "  ")
    rl.bold = True; rl.font.size = Pt(11); rl.font.color.rgb = NAVY; rl.font.name = 'Arial'
    rv = pp.add_run(text)
    rv.font.size = Pt(11); rv.font.color.rgb = DARK; rv.font.name = 'Arial'

doc.add_page_break()

# ── SECTION 1: EXECUTIVE SUMMARY ──────────────────────────────────

add_heading(doc, "Executive Summary")

make_highlight_box(doc,
    "Key Finding: A Supply Gap with Proven Demand",
    "The Lorton/Fairfax Station corridor sits at the center of one of the wealthiest, most golf-active regions "
    "in the United States — yet contains zero indoor golf simulator facilities. All existing competitors cluster "
    "10+ miles to the north toward the DC urban core, leaving an affluent suburban population of 300,000+ with "
    "no convenient indoor golf option. Workhouse Arts Center positions Fairway Golf Club to capture this untapped "
    "demand from day one."
)

add_spacer(doc, 10)

# Stat row table
stat_table = doc.add_table(rows=1, cols=4)
stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
stat_widths = [1.65, 1.65, 1.65, 1.65]
for i, (num, lbl, sub) in enumerate([
    ("0", "Simulator Competitors", "within 10 mi south"),
    ("$130K+", "Median HHI", "Fairfax County*"),
    ("51,000+", "Fort Belvoir", "employees within 5 mi"),
    ("5+", "Golf Courses", "within 7 miles"),
]):
    cell = stat_table.cell(0, i)
    set_cell_bg(cell, BG_GREEN)
    set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
    # clear default para
    cell.paragraphs[0].clear()
    p_num = cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_num.paragraph_format.space_before = Pt(4)
    p_num.paragraph_format.space_after = Pt(0)
    r_num = p_num.add_run(num)
    r_num.bold = True; r_num.font.size = Pt(24); r_num.font.color.rgb = GREEN; r_num.font.name = 'Arial'

    p_lbl = cell.add_paragraph()
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lbl.paragraph_format.space_before = Pt(0); p_lbl.paragraph_format.space_after = Pt(0)
    r_lbl = p_lbl.add_run(lbl)
    r_lbl.bold = True; r_lbl.font.size = Pt(9); r_lbl.font.color.rgb = NAVY; r_lbl.font.name = 'Arial'

    p_sub = cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0); p_sub.paragraph_format.space_after = Pt(4)
    r_sub = p_sub.add_run(sub)
    r_sub.italic = True; r_sub.font.size = Pt(8); r_sub.font.color.rgb = GRAY; r_sub.font.name = 'Arial'

# manually set col widths
tbl_el = stat_table._tbl
for row in stat_table.rows:
    for idx, cell in enumerate(row.cells):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(stat_widths[idx] * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

add_caveat(doc, "* Median HHI for Fairfax County derived from ACS 2022 5-year estimates; verify independently with current Census data")
add_spacer(doc, 6)

add_body(doc,
    "This report analyzes the local market opportunity for Fairway Golf Club, a proposed AI-powered indoor "
    "golf improvement club at Workhouse Arts Center, Lorton, Virginia. The analysis covers five areas: "
    "(1) demographics and purchasing power of the target corridor, (2) the competitive landscape and drive-time "
    "proximity of existing indoor golf facilities, (3) nearby golf course activity as a demand proxy, "
    "(4) national and regional indoor golf market growth, and (5) the Fort Belvoir military community as a "
    "supplemental customer segment."
)
add_body(doc,
    "Data sources include the U.S. Census Bureau (2020 Census, 2022 American Community Survey), Wikipedia, "
    "and industry knowledge current as of mid-2026. Market-sizing figures and competitor pricing are estimates; "
    "investors should independently verify before making decisions.",
    italic=True, size=9.5, color=GRAY
)

doc.add_page_break()

# ── SECTION 2: DEMOGRAPHICS ────────────────────────────────────────

add_heading(doc, "Section 1: Target Market Demographics")
add_heading(doc, "1.1 Primary Trade Area", level=2)

add_body(doc,
    "Fairway Golf Club's primary trade area is a 10-15 mile radius centered on Workhouse Arts Center, "
    "covering six distinct communities with strong household income and established golf culture."
)
add_spacer(doc, 6)

# Demographics table
demo_cols = ["Community", "Est. Population", "Approx. Median HHI", "Drive to Workhouse"]
demo_data = [
    ("Lorton / Fairfax Station",    "~55,000",    "$120K–$150K*",   "0–8 min",    BG_GREEN),
    ("Springfield",                  "~35,000",    "$95K–$115K*",    "10–15 min",  "FFFFFF"),
    ("Woodbridge",                   "~75,000",    "$85K–$100K*",    "15–20 min",  BG_GREEN),
    ("Dale City",                    "~75,000",    "$80K–$100K*",    "15–20 min",  "FFFFFF"),
    ("Fort Belvoir CDP (on-base)",   "7,637",      "$97,290 (Census)","5–8 min",   BG_GREEN),
    ("Newington / Kingstowne",       "~40,000",    "$110K–$130K*",   "10–15 min",  "FFFFFF"),
]

demo_table = doc.add_table(rows=len(demo_data)+1, cols=4)
demo_table.alignment = WD_TABLE_ALIGNMENT.LEFT
demo_col_w = [2.2, 1.3, 1.6, 1.4]

# Header
hdr = demo_table.rows[0]
for i, label in enumerate(demo_cols):
    c = hdr.cells[i]
    set_cell_bg(c, BG_NAV)
    set_cell_borders(c)
    set_cell_margins(c)
    cell_paragraph(c, label, bold=True, size=9.5, color=WHITE)

# Data rows
for ri, (community, pop, hhi, drive, bg) in enumerate(demo_data):
    row = demo_table.rows[ri+1]
    for ci, val in enumerate([community, pop, hhi, drive]):
        c = row.cells[ci]
        set_cell_bg(c, bg)
        set_cell_borders(c)
        set_cell_margins(c)
        is_comm = (ci == 0)
        cell_paragraph(c, val, bold=is_comm, size=9.5, color=NAVY if is_comm else DARK)

# set col widths
for row in demo_table.rows:
    for ci, cell in enumerate(row.cells):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(demo_col_w[ci] * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

add_caveat(doc, "* HHI estimates from ACS 2022 5-year estimates; verify exact figures via Census data for target census tracts")
add_spacer(doc, 6)

add_body(doc,
    "The combined addressable population within a 15-mile radius exceeds 300,000 residents. At the national "
    "golf participation rate of ~9-11% for adults, this suggests a potential golfer population of 25,000–35,000 "
    "within the primary trade area — more than sufficient to sustain a 2-bay operation and scale to 4 bays."
)

add_heading(doc, "1.2 Fairfax County: One of the Wealthiest Counties in America", level=2)
add_body(doc,
    "Fairfax County consistently ranks among the top 3-5 wealthiest counties in the United States by median "
    "household income. Key ACS-derived indicators:"
)
add_bullet(doc, "Median household income: approximately $130,000–$140,000 annually — one of the highest of any county in the US*")
add_bullet(doc, "Educational attainment: 60%+ of adults hold a bachelor's degree or higher*")
add_bullet(doc, "Employment base: federal government, defense contracting, technology (Amazon HQ2, Microsoft, Booz Allen, Leidos, SAIC), and professional services")
add_bullet(doc, "Population: 1.15 million+, the most populous jurisdiction in Virginia*")
add_caveat(doc, "* Verify with current ACS 5-year estimates; figures reflect knowledge as of mid-2026")
add_spacer(doc, 4)
add_body(doc, "This profile directly aligns with golf's core consumer: educated, professional, household income $75K+, ages 30–55. Fairfax County overrepresents this segment relative to national averages.")

add_heading(doc, "1.3 Lorton / Fairfax Station Micro-Market", level=2)
add_body(doc, "The immediate vicinity of Workhouse Arts Center has characteristics that particularly favor a premium entertainment concept:")
add_bullet(doc, "Workhouse Arts Center provides an existing visitor base, historic arts district credibility, and unique campus atmosphere")
add_bullet(doc, "Bunny Man Brewing is co-located on campus — natural food-and-beverage ecosystem, shared foot traffic")
add_bullet(doc, "Adjacent to Laurel Hill Golf Club, placing Fairway at the heart of existing outdoor golf demand")
add_bullet(doc, "I-95/Fairfax County Parkway interchange nearby — accessible from Fairfax County commuters and Prince William County residents")
add_bullet(doc, "No comparable premium entertainment venue exists in this specific corridor today")

doc.add_page_break()

# ── SECTION 3: COMPETITOR LANDSCAPE ───────────────────────────────

add_heading(doc, "Section 2: Competitive Landscape")
add_heading(doc, "2.1 The Geographic Supply Gap", level=2)

make_highlight_box(doc,
    "Critical Observation:",
    "Analysis of Google Maps search results for 'golf simulator' in the Lorton area confirms: all indoor golf "
    "simulator facilities in Northern Virginia are concentrated 10+ miles north of Workhouse, clustered toward the "
    "DC urban core and inner suburban ring. The area south of Workhouse — including Woodbridge, Dale City, Manassas, "
    "and the Route 1 / I-95 corridor into Prince William County — contains ZERO indoor golf simulators.",
    bg="FDF6E8", border_color="C9963B"
)
add_spacer(doc, 10)

# Competitor table
comp_cols = ["Facility", "Location", "Drive from Lorton*", "Offering", "Est. Pricing"]
comp_data = [
    ("CAFDExGO Golf",      "Springfield, VA",   "~10–15 min",  "Full Swing simulators, rental bays, lessons",             "~$45–60/hr**",          BG_GREEN),
    ("Uni Indoor Golf",    "Annandale, VA",      "~20–25 min",  "Basic simulator rental, minimal amenities",               "~$30/hr (budget)†",     "FFFFFF"),
    ("GOLFTEC Fairfax",    "Fairfax, VA",        "~20–25 min",  "Lesson-focused, TECFIT — not open sim rental",            "$100–150/hr (lessons)**",BG_GREEN),
    ("GOLFTEC Alexandria", "Alexandria, VA",     "~20–25 min",  "Same GOLFTEC model — lessons and swing analysis",         "$100–150/hr (lessons)**","FFFFFF"),
    ("Five Iron Golf",     "Washington, DC",     "~35–45 min",  "Premium urban: simulators, full bar, lounge, F&B",        "$50–85/hr (peak)**",    BG_GREEN),
    ("ParCiti Golf",       "Arlington, VA",      "~35–40 min",  "Premium urban: rooftop simulators, cocktail bar",         "$65–90/hr (est.)**",    "FFFFFF"),
]

comp_table = doc.add_table(rows=len(comp_data)+1, cols=5)
comp_table.alignment = WD_TABLE_ALIGNMENT.LEFT
comp_col_w = [1.4, 1.1, 1.1, 2.2, 1.3]

hdr = comp_table.rows[0]
for i, label in enumerate(comp_cols):
    c = hdr.cells[i]
    set_cell_bg(c, BG_NAV)
    set_cell_borders(c)
    set_cell_margins(c)
    cell_paragraph(c, label, bold=True, size=9, color=WHITE)

for ri, (name, loc, drive, offering, price, bg) in enumerate(comp_data):
    row = comp_table.rows[ri+1]
    vals = [name, loc, drive, offering, price]
    for ci, val in enumerate(vals):
        c = row.cells[ci]
        set_cell_bg(c, bg)
        set_cell_borders(c)
        set_cell_margins(c)
        cell_paragraph(c, val, bold=(ci==0), size=9, color=NAVY if ci==0 else DARK)

for row in comp_table.rows:
    for ci, cell in enumerate(row.cells):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(comp_col_w[ci] * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

add_caveat(doc, "* Drive time estimates based on typical off-peak traffic; I-95 peak hours can add 15-30 min")
add_caveat(doc, "** Pricing estimates from publicly available information; verify directly with each venue")
add_caveat(doc, "† Uni Indoor Golf (Annandale) visited by operator in July 2026: $30/hr confirmed, noted as 'budget, dark, tracking inconsistent'")

add_spacer(doc, 8)
add_heading(doc, "2.2 Competitive Positioning", level=2)
add_body(doc, "Fairway Golf Club occupies the suburban premium position — between budget simulators and expensive urban clubs:")

# Positioning table
pos_cols = ["Dimension", "Budget Competitors", "Urban Premium", "Fairway Golf Club"]
pos_data = [
    ("Price / hr",        "$25–35",                           "$65–90",                    "$55–70 (members ~$40)"),
    ("Location",          "Strip mall / industrial",          "Urban DC / Arlington",      "Arts campus, suburban community"),
    ("Technology",        "Basic simulator rental",           "Premium sim + F&B",         "AI coaching + Golfer360 profile"),
    ("Drive from Lorton", "20–25 min",                        "35–45 min",                 "0 min (on-site)"),
    ("Improvement focus", "None",                             "Some (Five Iron: lessons)",  "Core mission: every visit = better golfer"),
]

pos_table = doc.add_table(rows=len(pos_data)+1, cols=4)
pos_table.alignment = WD_TABLE_ALIGNMENT.LEFT
pos_col_w = [1.8, 1.8, 1.8, 2.2]

hdr = pos_table.rows[0]
for i, (label, bg) in enumerate(zip(pos_cols, [BG_NAV, BG_NAV, BG_NAV, BG_GRN])):
    c = hdr.cells[i]
    set_cell_bg(c, bg)
    set_cell_borders(c)
    set_cell_margins(c)
    cell_paragraph(c, label, bold=True, size=9.5, color=WHITE)

for ri, (dim, budget, urban, fairway) in enumerate(pos_data):
    row = pos_table.rows[ri+1]
    bg = BG_GREEN if ri % 2 == 0 else "FFFFFF"
    fg_bg = BG_GREEN
    for ci, (val, cell_bg) in enumerate(zip([dim, budget, urban, fairway], [bg, bg, bg, fg_bg])):
        c = row.cells[ci]
        set_cell_bg(c, cell_bg)
        set_cell_borders(c)
        set_cell_margins(c)
        is_bold = (ci == 0 or ci == 3)
        cell_paragraph(c, val, bold=is_bold, size=9.5, color=NAVY if is_bold else DARK)

for row in pos_table.rows:
    for ci, cell in enumerate(row.cells):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(pos_col_w[ci] * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

add_spacer(doc, 8)
add_body(doc,
    "Fairway Golf Club is the only option combining suburban convenience, premium experience, and technology-driven "
    "improvement for the Lorton corridor golfer. Urban premium clubs require 35-45 minutes (or more in DC traffic). "
    "Budget options offer no improvement pathway and a poor experience. Fairway fills this gap uniquely."
)

doc.add_page_break()

# ── SECTION 4: GOLF DEMAND INDICATORS ─────────────────────────────

add_heading(doc, "Section 3: Golf Demand Indicators — Nearby Courses")
add_body(doc,
    "Active outdoor golf courses within a facility's trade area are the strongest proxy for indoor golf demand. "
    "Indoor simulators thrive when serving existing golfers seeking year-round practice and improvement. "
    "Lorton scores exceptionally well on this metric."
)
add_spacer(doc, 8)

# Golf courses table
gc_cols = ["Golf Course", "Distance", "Type", "Character & Relevance"]
gc_data = [
    ("Laurel Hill Golf Club",        "<0.5 mi", "Public, 18-hole", "NVRPA-managed, highly regarded course on former Lorton prison grounds. Literally adjacent to Workhouse — any Laurel Hill golfer is a potential Fairway member.", BG_GREEN),
    ("Pohick Bay Golf Course",       "~3 mi",   "Public, 18-hole", "Pohick Bay Regional Park, wooded parkland setting. Family-friendly pricing attracts recreational golfers ideal for Fairway's beginner/improver segment.", "FFFFFF"),
    ("Fort Belvoir Golf Course",     "~3 mi",   "Military/semi-public", "On-base 18-hole course serving active duty, retirees, and DoD civilians. Direct pipeline to highly motivated golfers.", BG_GREEN),
    ("Old Hickory Golf Club",        "~5 mi",   "Public, 18-hole", "Well-reviewed Woodbridge course with active league programs. League players are prime targets for indoor league nights.", "FFFFFF"),
    ("Lake Ridge Golf Course",       "~6 mi",   "Public, 18-hole", "Woodbridge area, serves the growing Prince William County population. Demonstrates broad regional golf demand.", BG_GREEN),
    ("Burke Lake Golf Center",       "~7 mi",   "Par-3 / driving range", "NVRPA facility; driving range + par-3. Practice-oriented golfers and beginners — exactly the Fairway AI-coaching target market.", "FFFFFF"),
]

gc_table = doc.add_table(rows=len(gc_data)+1, cols=4)
gc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
gc_col_w = [1.8, 0.8, 1.3, 3.3]

hdr = gc_table.rows[0]
for i, label in enumerate(gc_cols):
    c = hdr.cells[i]
    set_cell_bg(c, BG_NAV)
    set_cell_borders(c)
    set_cell_margins(c)
    cell_paragraph(c, label, bold=True, size=9.5, color=WHITE)

for ri, (name, dist, tp, char, bg) in enumerate(gc_data):
    row = gc_table.rows[ri+1]
    for ci, val in enumerate([name, dist, tp, char]):
        c = row.cells[ci]
        set_cell_bg(c, bg)
        set_cell_borders(c)
        set_cell_margins(c)
        cell_paragraph(c, val, bold=(ci==0), size=9, color=NAVY if ci==0 else DARK)

for row in gc_table.rows:
    for ci, cell in enumerate(row.cells):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(gc_col_w[ci] * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

add_spacer(doc, 8)
add_heading(doc, "3.1 What This Means for Fairway Golf Club", level=2)
add_body(doc, "Five public golf courses plus an active military course within 7 miles creates a proven golfer population with nowhere convenient to practice indoors:")
add_bullet(doc, "Seasonal displacement: Virginia golfers lose 3-4 months of comfortable outdoor play per year to cold, rain, and shortened daylight. The nearest indoor option today is 20+ minutes away.")
add_bullet(doc, "Practice-to-play pipeline: Course golfers motivated to improve will pay for simulator time between rounds — especially those playing the demanding Laurel Hill layout literally next door.")
add_bullet(doc, "Social and league carryover: Courses like Old Hickory and Pohick Bay have active leagues. Fairway can recruit these groups for indoor league nights in the off-season and winter months.")
add_bullet(doc, "Beginner conversion: Burke Lake Golf Center attracts beginners using the driving range. Fairway's AI coaching platform specifically appeals to newer golfers who want structured, measurable improvement.")

doc.add_page_break()

# ── SECTION 5: MARKET SIZE ─────────────────────────────────────────

add_heading(doc, "Section 4: Indoor Golf Market Size & Revenue Model")
add_heading(doc, "4.1 National Market Context", level=2)

add_body(doc,
    "The indoor golf simulator industry has experienced exceptional growth since 2020, driven by three converging "
    "forces: post-pandemic demand for private, small-group social activities; major simulator technology improvements "
    "(Trackman, Full Swing, Foresight) making simulator play indistinguishable from on-course stats; and golf's "
    "broader participation surge from COVID-era outdoor activity adoption."
)
add_spacer(doc, 6)

# Two-col market table
mkt_table = doc.add_table(rows=1, cols=2)
mkt_table.alignment = WD_TABLE_ALIGNMENT.LEFT

left = mkt_table.cell(0, 0)
right = mkt_table.cell(0, 1)
for c, bg in [(left, BG_GREEN), (right, BG_GOLD)]:
    set_cell_bg(c, bg)
    set_cell_borders(c)
    set_cell_margins(c, top=120, bottom=120, left=140, right=140)

# set widths
for ci, cell in enumerate([left, right]):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(3.3 * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

cell_paragraph(left, "National Growth Indicators", bold=True, size=11, color=NAVY)
for item in [
    "• Estimated CAGR 2020–2025: 15–20%*",
    "• US indoor golf market estimated $1.0B+ by 2025–2026*",
    "• 41M+ Americans played golf in 2023 (rounds + off-course)*",
    "• Simulator technology now standard for serious golfer improvement",
]:
    add_cell_para(left, item, size=9.5, color=DARK)

cell_paragraph(right, "DC Metro Specifics", bold=True, size=11, color=NAVY)
for item in [
    "• High-income suburban market historically outperforms national averages for leisure spending",
    "• Five Iron Golf (DC), ParCiti (Arlington): proof of strong premium urban demand",
    "• Southern NoVA (Lorton corridor) currently unserved despite high suburban density*",
    "• Federal/DoD workforce culture aligns strongly with golf participation",
]:
    add_cell_para(right, item, size=9.5, color=DARK)

add_caveat(doc, "* Verify with current industry research: National Golf Course Owners Association, Golf Business Network, IBISWorld")
add_spacer(doc, 8)

add_heading(doc, "4.2 Unit Economics — Revenue Per Bay Per Year", level=2)
add_body(doc, "Based on DC Metro pricing and industry benchmarks, the following revenue scenarios apply to a premium suburban simulator bay:")
add_spacer(doc, 6)

rev_cols = ["Revenue Scenario", "Bay Utilization", "Avg. Blended Rate", "Annual Bay Revenue"]
rev_data = [
    ("Conservative (40% avg utilization)", "40% of 12 hrs/day", "$50/hr",  "$87,600 / bay",   "FFFFFF"),
    ("Moderate (55% avg utilization)",     "55% of 12 hrs/day", "$55/hr",  "$132,330 / bay",  BG_GREEN),
    ("Optimistic (70% avg utilization)",   "70% of 12 hrs/day", "$60/hr",  "$183,960 / bay",  "FFFFFF"),
]

rev_table = doc.add_table(rows=len(rev_data)+1, cols=4)
rev_table.alignment = WD_TABLE_ALIGNMENT.LEFT
rev_col_w = [2.6, 1.6, 1.4, 1.6]

hdr = rev_table.rows[0]
for i, label in enumerate(rev_cols):
    c = hdr.cells[i]
    set_cell_bg(c, BG_NAV)
    set_cell_borders(c)
    set_cell_margins(c)
    cell_paragraph(c, label, bold=True, size=9.5, color=WHITE)

for ri, (scenario, util, rate, revenue, bg) in enumerate(rev_data):
    row = rev_table.rows[ri+1]
    is_mid = (ri == 1)
    for ci, val in enumerate([scenario, util, rate, revenue]):
        c = row.cells[ci]
        set_cell_bg(c, bg)
        set_cell_borders(c)
        set_cell_margins(c)
        cell_paragraph(c, val, bold=(ci==3 and is_mid), size=9.5, color=GREEN if (ci==3 and is_mid) else DARK)

for row in rev_table.rows:
    for ci, cell in enumerate(row.cells):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(rev_col_w[ci] * 1440)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

add_caveat(doc, "* 12-hr operational day; actual utilization will be back-weighted to evenings and weekends; blended rate reflects mix of walk-in, member, and corporate bookings")
add_spacer(doc, 6)
add_body(doc,
    "At 2 bays and moderate utilization, Fairway Golf Club targets $250,000–$270,000 in annual bay rental revenue "
    "before membership premiums, AI coaching subscriptions, beverage sales, corporate events, and merchandise. "
    "Phase 1 financial modeling projects break-even at approximately 22% utilization — a conservative threshold "
    "suggesting meaningful downside protection."
)

doc.add_page_break()

# ── SECTION 6: FORT BELVOIR ────────────────────────────────────────

add_heading(doc, "Section 5: Fort Belvoir Military Community")
add_heading(doc, "5.1 Scale of the Opportunity", level=2)

add_body(doc,
    "Fort Belvoir, located approximately 5 miles from Workhouse Arts Center, is not a typical military installation. "
    "Following BRAC 2005 consolidations, it became one of the most significant DoD employment centers in the country "
    "— and one of the most compelling adjacent customer populations for Fairway Golf Club."
)
add_spacer(doc, 8)

# Fort Belvoir stats — two-col
fb_table = doc.add_table(rows=1, cols=2)
fb_table.alignment = WD_TABLE_ALIGNMENT.LEFT

left_fb = fb_table.cell(0, 0)
right_fb = fb_table.cell(0, 1)
set_cell_bg(left_fb, BG_GREEN)
set_cell_bg(right_fb, BG_GOLD)
for c in [left_fb, right_fb]:
    set_cell_borders(c)
    set_cell_margins(c, top=120, bottom=120, left=140, right=140)
    tc = c._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(3.3 * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

cell_paragraph(left_fb, "By the Numbers (2020 Census / Wikipedia)", bold=True, size=11, color=NAVY)
add_cell_para(left_fb, "51,000+", bold=True, size=22, color=GREEN)
add_cell_para(left_fb, "Total Fort Belvoir employees", size=9.5, color=DARK)
add_cell_para(left_fb, " ", size=6, color=WHITE)
add_cell_para(left_fb, "\"Nearly twice as many workers as the Pentagon\"", size=9.5, color=GRAY)
add_cell_para(left_fb, " ", size=6, color=WHITE)
add_cell_para(left_fb, "$97,290", bold=True, size=16, color=NAVY)
add_cell_para(left_fb, "Median household income (Fort Belvoir CDP, Census)", size=9.5, color=DARK)
add_cell_para(left_fb, "51.5% hold bachelor's degree or higher", size=9.5, color=DARK)

cell_paragraph(right_fb, "Who Works at Fort Belvoir", bold=True, size=11, color=NAVY)
for item in [
    "• Active duty Army (29th Infantry Division, plus 10 Army major commands)",
    "• DoD civilians and defense contractors",
    "• Defense Logistics Agency (headquarters)",
    "• National Geospatial-Intelligence Agency (NGA Campus East)",
    "• Missile Defense Agency, DTRA, DCAA, DAU",
    "• 26 Department of Defense agencies",
    "• 86.6% of CDP households are married couples (Census 2020)",
    "• Avg. household size: 4.21 people",
]:
    add_cell_para(right_fb, item, size=9.5, color=DARK)

add_spacer(doc, 10)
add_heading(doc, "5.2 Military Golf Culture & Segment Value", level=2)
add_body(doc, "Golf has historically been deeply embedded in military culture, particularly among officers and senior non-commissioned officers. Fort Belvoir-specific indicators:")
add_bullet(doc, "Woody's Golf Course (Fort Belvoir Golf Course): On-base 18-hole facility with active use by active duty, retirees, and DoD civilians — a directly accessible pipeline of motivated golfers who would benefit from Fairway's AI coaching.")
add_bullet(doc, "Cultural fit: The data-driven, metrics-focused nature of Fairway's Golfer360 platform mirrors the culture of military and DoD professionals — a population that responds strongly to measurable progress tracking.")
add_bullet(doc, "Corporate group bookings: Fort Belvoir agencies and units represent a significant events opportunity. Agency nights, unit socials, and contractor team-building outings could drive regular block bookings.")
add_bullet(doc, "MWR Partnership potential: A relationship with Fort Belvoir's Morale, Welfare and Recreation (MWR) program could direct active duty and retirees to Fairway as a preferred off-base recreational destination.")
add_spacer(doc, 4)
add_body(doc,
    "The 51,000-person Fort Belvoir workforce represents a sustained, predictable customer pipeline within a "
    "5-8 minute drive of Workhouse Arts Center — a population that skews toward disciplined, recreational-activity-oriented "
    "adults with above-average education and household income."
)

doc.add_page_break()

# ── SECTION 7: SWOT + INVESTMENT THESIS ───────────────────────────

add_heading(doc, "Section 6: Strategic Assessment")
add_heading(doc, "6.1 SWOT Analysis", level=2)
add_spacer(doc, 6)

# SWOT 2x2 table
swot_table = doc.add_table(rows=2, cols=2)
swot_table.alignment = WD_TABLE_ALIGNMENT.LEFT

swot_items = [
    ("STRENGTHS", "E8F5E9", RGBColor(0x2A, 0x60, 0x41), [
        "Zero direct competitors in the trade area — first-mover advantage in an untapped corridor",
        "Five golf courses within 7 miles = proven active golfer demand, ready to convert to indoor practice",
        "Workhouse Arts Center campus provides arts district credibility, existing foot traffic, brewery co-tenant",
        "Operator Salesforce/Agentforce expertise is a defensible technology moat competitors cannot easily replicate",
    ]),
    ("WEAKNESSES", "FFF8E1", RGBColor(0x8B, 0x69, 0x14), [
        "2-bay Phase 1 limits capacity during peak demand — Saturday evenings could have waitlists quickly",
        "New entrant brand in a market with no indoor golf precedent; awareness-building required",
        "Virginia ABC licensing for on-site beverage sales adds complexity and 6-12 month timeline risk",
    ]),
    ("OPPORTUNITIES", "E8F0FE", RGBColor(0x1A, 0x4A, 0x8A), [
        "Phase 2 expansion to 4 bays — low incremental capital cost once core buildout is complete",
        "Fort Belvoir MWR partnership could unlock thousands of military and retiree members",
        "Corporate events market: DoD agencies, defense contractors, law firms and consulting firms in Fairfax County",
        "Growing millennial/Gen Z golf participation builds the next generation of loyal Fairway members",
    ]),
    ("THREATS", "FEECEC", RGBColor(0x8B, 0x1A, 0x1A), [
        "National chains (Five Iron Golf, Topgolf Swing Suites) may enter this corridor once market is validated",
        "Consumer-grade simulator hardware continues to improve — at-home competition is a long-term risk",
        "Workhouse Arts Center lease negotiations: space availability and terms are not yet confirmed",
    ]),
]

for i, (title, bg, title_color, bullets) in enumerate(swot_items):
    row_i = i // 2
    col_i = i % 2
    cell = swot_table.cell(row_i, col_i)
    set_cell_bg(cell, bg)
    set_cell_borders(cell)
    set_cell_margins(cell, top=120, bottom=120, left=140, right=140)

    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(3.3 * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

    cell_paragraph(cell, title, bold=True, size=11, color=title_color)
    for b in bullets:
        add_cell_para(cell, f"• {b}", size=9.5, color=DARK, space_before=3)

add_spacer(doc, 12)

add_heading(doc, "6.2 Investment Thesis", level=2)

make_highlight_box(doc,
    "The Case for Fairway Golf Club",
    "The Lorton / Fairfax Station corridor presents a rare convergence of high purchasing power, active golf demand, "
    "and complete absence of indoor golf supply. Five public golf courses within 7 miles signal a large, active golfer "
    "population. Fort Belvoir's 51,000-person workforce sits 5 minutes away. Fairfax County's $130K+ median household "
    "income ensures the target customer can afford premium membership pricing. And every existing competitor is at least "
    "10 miles north. Fairway Golf Club does not need to create demand. It needs to be in the right place when existing "
    "demand looks for a home — and no one else is there yet.",
    bg=BG_GOLD, border_color="C9963B"
)

add_spacer(doc, 12)
add_heading(doc, "6.3 Data Verification Checklist", level=2)
add_body(doc, "The following items should be independently verified before finalizing investor materials:")
add_bullet(doc, "Current ACS 5-year estimates for Lorton, Fairfax Station, Woodbridge, and Dale City census tracts (Census.gov)")
add_bullet(doc, "Current pricing at CAFDExGO Golf, Uni Indoor Golf, Five Iron Golf DC, and ParCiti (call or visit each)")
add_bullet(doc, "Competitor status: confirm all listed facilities are currently open and active")
add_bullet(doc, "Laurel Hill Golf Club annual round count and peak-season demand data (call NVRPA)")
add_bullet(doc, "Fort Belvoir MWR golf participation and on-base population data (Fort Belvoir public affairs)")
add_bullet(doc, "Indoor golf industry market size reports (NGCOA, Golf Business Network, IBISWorld)")
add_bullet(doc, "Workhouse Arts Center available square footage, lease rate, and existing foot traffic data")

add_spacer(doc, 16)
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_foot.paragraph_format.space_before = Pt(8)
r_foot = p_foot.add_run("Prepared by Fairway Golf Club operating team  |  July 2026  |  CONFIDENTIAL — For investor use only")
r_foot.italic = True; r_foot.font.size = Pt(9); r_foot.font.color.rgb = GRAY; r_foot.font.name = 'Arial'

# Save
out_path = "Fairway_Golf_Club_Market_Analysis.docx"
doc.save(out_path)
print(f"Done: {out_path}")
